from collections import defaultdict
from contextvars import ContextVar
import os
import re
import markdown2

# from weasyprint import HTML, CSS
try:
    from weasyprint import HTML, CSS
except (ImportError, OSError):
    HTML = None
    CSS = None
    print("⚠️  WeasyPrint not available - PDF generation disabled")


import logging
from fontTools.misc.loggingTools import configLogger
import yaml
from dotenv import load_dotenv

load_dotenv()  # loads .env if present


# Silence noisy logs
configLogger(level=logging.CRITICAL)
logging.getLogger("fontTools").setLevel(logging.CRITICAL)
logging.getLogger("weasyprint").setLevel(logging.CRITICAL)
logging.getLogger().setLevel(logging.CRITICAL)

# -------------------------------------------------------------------
# GLOBAL TOKEN USAGE TRACKER
# -------------------------------------------------------------------
_token_usage_history_ctx = ContextVar("token_usage_history", default=None)


class TokenUsageHistory:
    """Context-local token tracker to keep concurrent runs isolated."""

    def _current(self):
        history = _token_usage_history_ctx.get()
        if history is None:
            history = []
            _token_usage_history_ctx.set(history)
        return history

    def append(self, item):
        self._current().append(item)

    def clear(self):
        self._current().clear()

    def reset(self):
        _token_usage_history_ctx.set([])

    def __iter__(self):
        return iter(self._current())

    def __len__(self):
        return len(self._current())

    def __getitem__(self, index):
        return self._current()[index]


token_usage_history = TokenUsageHistory()

# -------------------------------------------------------------------
# AI CORE CREDENTIAL LOADER (FINAL, CLEAN VERSION)
# -------------------------------------------------------------------

import json


def _load_aicore_credentials():
    """Load AI Core credentials from VCAP_SERVICES or environment variables."""
    # --- 1) Try Cloud Foundry VCAP ---
    if "VCAP_SERVICES" in os.environ:
        try:
            vcap = json.loads(os.environ["VCAP_SERVICES"])
            aicore = vcap.get("aicore", [])

            if aicore:
                creds = aicore[0]["credentials"]
                return {
                    "AICORE_CLIENT_ID": creds.get("clientid"),
                    "AICORE_CLIENT_SECRET": creds.get("clientsecret"),
                    "AICORE_AUTH_URL": creds.get("url"),
                    "AICORE_BASE_URL": creds["serviceurls"]["AI_API_URL"] + "/v2",
                    "AICORE_RESOURCE_GROUP": os.getenv(
                        "AICORE_RESOURCE_GROUP", "default"
                    ),
                }
        except Exception as e:
            print("⚠️ Failed to load AI Core from VCAP:", e)

    # --- 2) Fallback to environment variables (.env for local) ---
    return {
        "AICORE_CLIENT_ID": os.getenv("AICORE_CLIENT_ID"),
        "AICORE_CLIENT_SECRET": os.getenv("AICORE_CLIENT_SECRET"),
        "AICORE_AUTH_URL": os.getenv("AICORE_AUTH_URL"),
        "AICORE_BASE_URL": os.getenv("AICORE_BASE_URL"),
        "AICORE_RESOURCE_GROUP": os.getenv("AICORE_RESOURCE_GROUP", "default"),
    }


def _set_env_if_valid(key, value):
    """Set an environment variable only if it has a valid value."""
    if value and value != "None":
        os.environ[key] = value
    else:
        print(f"⚠️ WARNING: {key} is missing or empty, not setting env var.")


# Load credentials ONCE at module import (skip if DEBUG mode)
debug_mode = os.getenv("DEBUG", "false").lower() == "true"

if not debug_mode:
    creds = _load_aicore_credentials()

    _set_env_if_valid("AICORE_CLIENT_ID", creds.get("AICORE_CLIENT_ID"))
    _set_env_if_valid("AICORE_CLIENT_SECRET", creds.get("AICORE_CLIENT_SECRET"))
    _set_env_if_valid("AICORE_AUTH_URL", creds.get("AICORE_AUTH_URL"))
    _set_env_if_valid("AICORE_BASE_URL", creds.get("AICORE_BASE_URL"))
    _set_env_if_valid("AICORE_RESOURCE_GROUP", creds.get("AICORE_RESOURCE_GROUP"))

    print("🔐 AI Core Environment Initialized")
    # print("   AICORE_BASE_URL =", os.environ.get("AICORE_BASE_URL"))
else:
    print("🔧 DEBUG MODE: Skipping AI Core credential loading")

# -------------------------------------------------------------------
# IMPORT GEN_AI_HUB ONLY IF NOT IN DEBUG MODE
# -------------------------------------------------------------------

# Delay import to avoid loading SAP AI Core in DEBUG mode
chat = None
if os.getenv("DEBUG", "false").lower() != "true":
    try:
        from gen_ai_hub.proxy.native.openai import chat
    except Exception as e:
        print(f"⚠️ WARNING: Could not import gen_ai_hub: {e}")
        print("⚠️ If using DEBUG mode, this is expected and can be ignored.")


# -------------------------------------------------------------------
# PROMPT & WORKFLOW UTILITIES
# -------------------------------------------------------------------


def load_prompt(agent_name: str):
    path = os.path.join("prompts", f"{agent_name}.yaml")
    with open(path, "r", encoding="utf-8") as f:
        return yaml.safe_load(f)


def format_user_prompt(prompt_def, state):
    params = {key: state[key] for key in prompt_def.get("parameters", {})}
    return prompt_def["user_prompt"].format(**params)


def parse_route(content):
    """Extract [ROUTE:*] directive."""
    if isinstance(content, list):
        content = " ".join(str(c.get("text", c)) for c in content)

    match = re.search(r"\[ROUTE:\s*(\w+)\]", str(content))
    return match.group(1) if match else "output_reviewer"


# -------------------------------------------------------------------
# TOKEN USAGE HANDLING
# -------------------------------------------------------------------


def add_token_usage(response, agent_name):
    """
    Extract token usage from LLM response and add to history.

    Works with multiple response formats:
    - LangChain ChatOpenAI/ChatVertexAI responses (usage_metadata attribute)
    - Direct API responses (usage attribute)

    Response format from SAP AI Core via LangChain:
    - response.usage_metadata.input_tokens (prompt tokens)
    - response.usage_metadata.output_tokens (completion tokens)
    - response.usage_metadata.total_tokens
    """

    # Try LangChain format first (usage_metadata)
    usage_metadata = getattr(response, "usage_metadata", None)

    if usage_metadata and isinstance(usage_metadata, dict):
        # LangChain format - usage_metadata is a dict
        usage_dict = {
            "input_tokens": usage_metadata.get("input_tokens", 0),
            "output_tokens": usage_metadata.get("output_tokens", 0),
            "total_tokens": usage_metadata.get("total_tokens", 0),
        }
        token_usage_history.append({"agent": agent_name, "usage": usage_dict})
        return

    # Try direct API format (usage)
    usage = getattr(response, "usage", None)

    if usage:
        # Direct API format (OpenAI response)
        usage_dict = {
            "input_tokens": getattr(usage, "prompt_tokens", 0),
            "output_tokens": getattr(usage, "completion_tokens", 0),
            "total_tokens": getattr(usage, "total_tokens", 0),
        }
        token_usage_history.append({"agent": agent_name, "usage": usage_dict})
        return

    # No usage data found - log warning only in debug mode
    if os.getenv("DEBUG", "False").lower() == "true":
        print(f"⚠️ No usage data found for {agent_name}")


def print_total_token_usage(token_usage_history):
    """Print summary of token usage across all agents."""
    total_input = 0
    total_output = 0
    total_tokens = 0

    print("\n" + "=" * 70)
    print("📊 TOKEN USAGE SUMMARY")
    print("=" * 70)

    for entry in token_usage_history:
        agent = entry.get("agent", "unknown")
        usage = entry.get("usage", {})

        input_tokens = usage.get("input_tokens", 0)
        output_tokens = usage.get("output_tokens", 0)
        total = usage.get("total_tokens", input_tokens + output_tokens)

        total_input += input_tokens
        total_output += output_tokens
        total_tokens += total

        print(
            f"  {agent:30s} | Input: {input_tokens:6d} | Output: {output_tokens:6d} | Total: {total:6d}"
        )

    print("-" * 70)
    print(
        f"  {'TOTAL':30s} | Input: {total_input:6d} | Output: {total_output:6d} | Total: {total_tokens:6d}"
    )
    print("=" * 70)


# -------------------------------------------------------------------
# LOGGING WRAPPER FOR LANGGRAPH
# -------------------------------------------------------------------


def logging_wrapper(fn, name):
    def wrapped(state):
        print(f"\n🟦 Running node: {name}")
        before = len(token_usage_history)
        result = fn(state)
        # print(f"🟩 Output from {name}: {result}")

        new_usage = token_usage_history[before:]
        if new_usage:
            print(f"🟨 Token usage for {name}: {new_usage}")

        return result

    return wrapped


# -------------------------------------------------------------------
# MARKDOWN CLEANING + PDF GENERATION (UNCHANGED)
# -------------------------------------------------------------------


def clean_markdown(md: str) -> str:
    lines = md.splitlines()
    cleaned = []
    removed = []
    started = False

    ai_re = re.compile(
        r"(as the project manager|as an ai|as a ai|assistant:|system:|human:|"
        r"consolidated specification|final, consolidated|facilitated discussions)",
        re.IGNORECASE,
    )

    for line in lines:
        l = line.strip()

        if l.lower().startswith("[route:"):
            removed.append(line)
            continue

        if started:
            cleaned.append(line)
            continue

        if l.startswith("#") or l.startswith("**Page") or l.startswith("|"):
            started = True
            cleaned.append(line)
            continue

        if ai_re.search(l):
            removed.append(line)
            continue

        if not l:
            removed.append(line)
            continue

        started = True
        cleaned.append(line)

    return "\n".join(cleaned) + "\n"


def convert_markdown_to_pdf(
    md_file: str, output_pdf: str, page_layout: str = "A4_original"
):
    """
    Convert Markdown to PDF with Mermaid diagram support and flexible page layouts.

    Args:
        md_file: Path to input Markdown file
        output_pdf: Path to output PDF file
        page_layout: Page layout mode - "A4_original", "A3_landscape", or "A4_A3_mixed"
            - "A4_original": A4 portrait with simple styling (DEFAULT, matches classic behavior)
            - "A3_landscape": All pages in A3 landscape (best for wide tables/diagrams)
            - "A4_A3_mixed": A4 portrait for text, A3 landscape for wide content

    Returns:
        Path to generated PDF file
    """
    import re
    import html
    import base64
    import zlib
    import urllib.request
    import urllib.error

    # Read markdown content
    with open(md_file, "r", encoding="utf-8") as f:
        markdown_content = f.read()

    # Remove any leaked placeholder tokens from LLM output (safety guard)
    markdown_content = re.sub(r"\{\{MERMAID_IMG_\d+\}\}", "", markdown_content)

    # Step 1: Extract and convert Mermaid diagrams to PNG
    def _fetch_url_bytes(url: str):
        """Fetch URL bytes with browser-like headers to reduce 403 responses."""
        request = urllib.request.Request(
            url,
            headers={
                "User-Agent": "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 "
                "(KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36",
                "Accept": "image/png,image/svg+xml,image/*;q=0.8,*/*;q=0.5",
            },
        )
        with urllib.request.urlopen(request, timeout=60) as response:
            return response.read()

    def _post_url_bytes(url: str, body: bytes, content_type: str):
        """POST bytes and return response payload."""
        request = urllib.request.Request(
            url,
            data=body,
            method="POST",
            headers={
                "User-Agent": "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 "
                "(KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36",
                "Accept": "image/png,image/*;q=0.8,*/*;q=0.5",
                "Content-Type": content_type,
            },
        )
        with urllib.request.urlopen(request, timeout=60) as response:
            return response.read()

    def _sanitize_mermaid_code(mermaid_code: str) -> str:
        """Normalize Mermaid text for better renderer compatibility."""
        code = (mermaid_code or "").replace("\r\n", "\n").replace("\r", "\n")
        code = code.replace("\ufeff", "").replace("\u200b", "")
        code = code.replace("“", '"').replace("”", '"').replace("’", "'")
        code = code.strip()

        # Remove accidental fence lines if present
        code = re.sub(
            r"^\s*```(?:mermaid)?\s*$", "", code, flags=re.IGNORECASE | re.MULTILINE
        )
        code = re.sub(r"^\s*```\s*$", "", code, flags=re.MULTILINE)

        lines = [line.rstrip() for line in code.split("\n")]
        lines = [
            line
            for line in lines
            if line.strip() and not re.search(r"\{\{MERMAID_IMG_\d+\}\}", line)
        ]
        if not lines:
            return ""

        first = lines[0].strip()
        valid_prefixes = (
            "flowchart",
            "graph",
            "sequenceDiagram",
            "classDiagram",
            "stateDiagram",
            "erDiagram",
            "journey",
            "gantt",
            "pie",
            "mindmap",
            "timeline",
            "quadrantChart",
            "requirementDiagram",
            "gitGraph",
            "C4Context",
            "C4Container",
            "C4Component",
            "C4Dynamic",
            "C4Deployment",
            "sankey-beta",
            "xychart-beta",
            "block-beta",
        )
        if not first.startswith(valid_prefixes):
            lines.insert(0, "flowchart TD")

        return "\n".join(lines).strip()

    def _repair_mermaid_code(mermaid_code: str) -> str:
        """Best-effort repair for common LLM Mermaid syntax issues."""
        repaired_lines = []
        for line in mermaid_code.split("\n"):
            stripped = line.strip()

            # Repair dangling flow edges like 'A-->'
            m = re.match(r"^([A-Za-z0-9_]+)\s*--?>\s*$", stripped)
            if m:
                repaired_lines.append(f"{m.group(1)}-->UNRESOLVED_NODE")
                continue

            # Repair dangling labeled edges like 'A-->|Yes|'
            m = re.match(r"^([A-Za-z0-9_]+)\s*--?>\|([^|]+)\|\s*$", stripped)
            if m:
                repaired_lines.append(f"{m.group(1)}-->|{m.group(2)}|UNRESOLVED_NODE")
                continue

            repaired_lines.append(line)

        repaired = "\n".join(repaired_lines)
        # Ensure placeholder node exists when used
        if "UNRESOLVED_NODE" in repaired and "UNRESOLVED_NODE[" not in repaired:
            repaired += "\nUNRESOLVED_NODE[Continue]"

        return repaired

    def _try_render_with_providers(mermaid_code: str):
        """Try Kroki then mermaid.ink. Returns (img_b64, error_text)."""
        errors = []

        # 1) Kroki POST text/plain (most reliable for long/complex diagrams)
        try:
            img_bytes = _post_url_bytes(
                "https://kroki.io/mermaid/png",
                mermaid_code.encode("utf-8"),
                "text/plain; charset=utf-8",
            )
            return base64.b64encode(img_bytes).decode("utf-8"), None
        except urllib.error.HTTPError as e:
            errors.append(f"Kroki POST text: HTTP {e.code}")
        except Exception as e:
            errors.append(f"Kroki POST text: {e}")

        # 2) Kroki POST JSON fallback
        try:
            payload = json.dumps(
                {
                    "diagram_source": mermaid_code,
                    "diagram_type": "mermaid",
                    "output_format": "png",
                }
            ).encode("utf-8")
            img_bytes = _post_url_bytes(
                "https://kroki.io/", payload, "application/json"
            )
            return base64.b64encode(img_bytes).decode("utf-8"), None
        except urllib.error.HTTPError as e:
            errors.append(f"Kroki POST json: HTTP {e.code}")
        except Exception as e:
            errors.append(f"Kroki POST json: {e}")

        # 3) Kroki GET encoded fallback
        try:
            compressed = zlib.compress(mermaid_code.encode("utf-8"), level=9)
            encoded = base64.urlsafe_b64encode(compressed).decode("utf-8")
            url = f"https://kroki.io/mermaid/png/{encoded}"
            img_bytes = _fetch_url_bytes(url)
            return base64.b64encode(img_bytes).decode("utf-8"), None
        except urllib.error.HTTPError as e:
            errors.append(f"Kroki GET: HTTP {e.code}")
        except Exception as e:
            errors.append(f"Kroki GET: {e}")

        # 4) Mermaid.ink fallback
        try:
            encoded_plain = base64.urlsafe_b64encode(
                mermaid_code.encode("utf-8")
            ).decode("utf-8")
            url = f"https://mermaid.ink/img/{encoded_plain}"
            img_bytes = _fetch_url_bytes(url)
            return base64.b64encode(img_bytes).decode("utf-8"), None
        except urllib.error.HTTPError as e:
            errors.append(f"mermaid.ink: HTTP {e.code}")
        except Exception as e:
            errors.append(f"mermaid.ink: {e}")

        return None, "; ".join(errors)

    def mermaid_to_png(mermaid_code):
        """Convert Mermaid code to PNG. Returns (base64_png, error_message)."""
        normalized_code = _sanitize_mermaid_code(mermaid_code)
        if not normalized_code:
            return None, "empty_mermaid_code"

        # Attempt 1: render sanitized source
        img_b64, err = _try_render_with_providers(normalized_code)
        if img_b64:
            return img_b64, None

        # Attempt 2: repair common syntax issues and retry
        repaired_code = _repair_mermaid_code(normalized_code)
        if repaired_code != normalized_code:
            img_b64, repaired_err = _try_render_with_providers(repaired_code)
            img_b64, fallback_err = _try_render_with_providers(fallback_visual)
            if img_b64:
                return img_b64, None
            err = f"{err}; repaired: {repaired_err}"

        # Attempt 3: guaranteed-valid visual fallback image
        fallback_visual = (
            "flowchart TD\n"
            "A[Diagram could not be rendered]\n"
            "B[See Mermaid source in document text]\n"
            "A-->B"
        )
        img_b64, fallback_err = _try_render_with_providers(fallback_visual)
        if img_b64:
            return img_b64, None
        err = f"{err}; visual-fallback: {fallback_err}"

        return None, err

    # Find all Mermaid code blocks
    mermaid_pattern = r"```mermaid[^\n\r]*\r?\n(.*?)```"
    mermaid_blocks = re.findall(mermaid_pattern, markdown_content, re.DOTALL)

    # Convert each Mermaid block to PNG
    mermaid_results = []
    failed_mermaid_details = []
    for mermaid_code in mermaid_blocks:
        img_b64, err = mermaid_to_png(mermaid_code)
        mermaid_results.append(
            {
                "code": mermaid_code,
                "img_b64": img_b64,
                "error": err,
            }
        )
        if err:
            failed_mermaid_details.append(err)

    if failed_mermaid_details:
        print(
            f"ℹ️ Mermaid rendering unavailable for {len(failed_mermaid_details)} diagram(s); "
            "embedded source fallback in PDF."
        )

    # Replace Mermaid blocks with image placeholders
    modified_markdown = markdown_content
    for i in range(len(mermaid_results)):
        placeholder = f"{{{{MERMAID_IMG_{i}}}}}"
        modified_markdown = re.sub(
            r"```mermaid[^\n\r]*\r?\n.*?```",
            placeholder,
            modified_markdown,
            count=1,
            flags=re.DOTALL,
        )

    # Step 2: Convert Markdown to HTML
    html_body = markdown2.markdown(
        modified_markdown,
        extras=["fenced-code-blocks", "tables", "strike", "code-friendly"],
    )

    # Step 3: Replace placeholders in HTML (supports inline placeholder usage)
    for i, result in enumerate(mermaid_results):
        placeholder = f"{{{{MERMAID_IMG_{i}}}}}"
        img_b64 = result["img_b64"]

        if img_b64:
            replacement_html = (
                f'<div class="mermaid-diagram">'
                f'<img src="data:image/png;base64,{img_b64}" alt="Diagram" />'
                f"</div>"
            )
        else:
            escaped_code = html.escape(result["code"])
            replacement_html = (
                '<div class="mermaid-diagram-fallback">'
                "<p><strong>Diagram source (rendering unavailable in runtime):</strong></p>"
                f"<pre>{escaped_code}</pre>"
                "</div>"
            )

        # Replace both standalone and inline placeholders
        html_body = html_body.replace(f"<p>{placeholder}</p>", replacement_html)
        html_body = html_body.replace(placeholder, replacement_html)

    # Final guard: remove any unresolved placeholders from output
    html_body = re.sub(r"\{\{MERMAID_IMG_\d+\}\}", "", html_body)

    # Step 4: Generate CSS based on page layout mode
    if page_layout == "A4_original":
        # A4 layout with proper table handling for SAP-style documentation.
        css = """
        @page {
            size: A4 portrait;
            margin: 0.6in 0.5in 0.6in 0.5in;
        }

        body {
            font-family: 'Segoe UI', Arial, Helvetica, sans-serif;
            font-size: 10pt;
            line-height: 1.5;
            color: #333;
        }

        h1 {
            font-size: 18pt;
            color: #1a1a2e;
            border-bottom: 3px solid #0070f3;
            padding-bottom: 6px;
            margin-top: 24px;
            margin-bottom: 12px;
            page-break-after: avoid;
        }

        h2 {
            font-size: 14pt;
            color: #1a1a2e;
            border-bottom: 1px solid #ccc;
            padding-bottom: 4px;
            margin-top: 20px;
            margin-bottom: 10px;
            page-break-after: avoid;
        }

        h3 {
            font-size: 12pt;
            color: #333;
            margin-top: 16px;
            margin-bottom: 8px;
            page-break-after: avoid;
        }

        h4 {
            font-size: 11pt;
            color: #444;
            margin-top: 12px;
            margin-bottom: 6px;
        }

        table {
            width: 100%;
            border-collapse: collapse;
            table-layout: fixed;
            margin: 10px 0;
            font-size: 8.5pt;
            word-wrap: break-word;
            overflow-wrap: break-word;
        }

        td, th {
            border: 1px solid #999;
            padding: 4px 6px;
            text-align: left;
            vertical-align: top;
            word-wrap: break-word;
            overflow-wrap: break-word;
            hyphens: auto;
        }

        th {
            background-color: #e8eef4;
            font-weight: bold;
            font-size: 8.5pt;
            color: #1a1a2e;
        }

        tr {
            page-break-inside: avoid;
        }

        code {
            background-color: #f4f4f4;
            padding: 1px 4px;
            border-radius: 3px;
            font-family: 'Courier New', monospace;
            font-size: 8pt;
        }

        pre {
            background-color: #f4f4f4;
            padding: 8px;
            border-radius: 4px;
            font-size: 8pt;
            white-space: pre-wrap;
            word-break: break-word;
            overflow-x: hidden;
            page-break-inside: avoid;
        }

        blockquote {
            border-left: 3px solid #f0ad4e;
            padding: 6px 12px;
            margin: 10px 0;
            background: #fff8e1;
            font-size: 9pt;
        }

        hr {
            border: none;
            border-top: 1px solid #ccc;
            margin: 16px 0;
        }

        ul, ol {
            padding-left: 20px;
            margin: 6px 0;
        }

        li {
            margin-bottom: 3px;
        }

        /* Mermaid diagram support */
        .mermaid-diagram {
            margin: 16px 0;
            text-align: center;
            page-break-inside: avoid;
        }

        .mermaid-diagram-fallback {
            margin: 16px 0;
            padding: 10px;
            border: 1px solid #ddd;
            background: #f9f9f9;
            page-break-inside: avoid;
        }

        .mermaid-diagram-fallback pre {
            white-space: pre-wrap;
            word-break: break-word;
            font-size: 8pt;
        }

        .mermaid-diagram img {
            max-width: 100%;
            height: auto;
        }
        """
    elif page_layout == "A4_A3_mixed":
        # Mixed layout: A4 portrait for text, A3 landscape for wide content
        css = """
        @page {
            size: A4 portrait;
            margin: 0.75in;
        }
        
        @page wide {
            size: A3 landscape;
            margin: 0.5in;
        }
        
        body {
            font-family: Arial, sans-serif;
            line-height: 1.6;
            font-size: 11pt;
        }
        
        h1 {
            font-size: 18pt;
            margin-top: 20px;
            margin-bottom: 15px;
            page-break-after: avoid;
        }
        
        h2 {
            font-size: 14pt;
            margin-top: 15px;
            margin-bottom: 10px;
            page-break-after: avoid;
        }
        
        /* Force wide tables to A3 landscape pages */
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 15px 0;
            font-size: 9pt;
        }
        
        /* Tables with >6 columns use A3 landscape */
        table:has(th:nth-child(7)) {
            page: wide;
            page-break-before: always;
            page-break-after: always;
        }
        
        td, th {
            border: 1px solid #999;
            padding: 5px 8px;
            text-align: left;
            vertical-align: top;
        }
        
        th {
            background-color: #e0e0e0;
            font-weight: bold;
        }
        
        /* Mermaid diagrams on A3 landscape */
        .mermaid-diagram {
            page: wide;
            page-break-before: always;
            page-break-after: always;
            margin: 20px 0;
            text-align: center;
        }

        .mermaid-diagram-fallback {
            page: wide;
            page-break-before: always;
            page-break-after: always;
            margin: 20px 0;
            padding: 10px;
            border: 1px solid #ddd;
            background: #f9f9f9;
        }

        .mermaid-diagram-fallback pre {
            white-space: pre-wrap;
            word-break: break-word;
            font-size: 9pt;
        }
        
        .mermaid-diagram img {
            max-width: 100%;
            height: auto;
            display: block;
            margin: 0 auto;
        }
        
        tr {
            page-break-inside: avoid;
        }
        """
    else:
        # A3 landscape layout: All content in A3 landscape (best for wide tables/diagrams)
        css = """
        @page {
            size: A3 landscape;
            margin: 0.5in;
        }
        
        body {
            font-family: Arial, sans-serif;
            line-height: 1.6;
            font-size: 10pt;
        }
        
        h1 {
            font-size: 18pt;
            margin-top: 20px;
            margin-bottom: 15px;
            page-break-after: avoid;
        }
        
        h2 {
            font-size: 14pt;
            margin-top: 15px;
            margin-bottom: 10px;
            page-break-after: avoid;
        }
        
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 15px 0;
            font-size: 9pt;
            page-break-inside: avoid;
        }
        
        td, th {
            border: 1px solid #999;
            padding: 5px 8px;
            text-align: left;
            vertical-align: top;
        }
        
        th {
            background-color: #e0e0e0;
            font-weight: bold;
        }
        
        code {
            background-color: #f4f4f4;
            padding: 2px 5px;
            border-radius: 3px;
            font-family: 'Courier New', monospace;
            font-size: 8pt;
        }
        
        pre {
            background-color: #f4f4f4;
            padding: 10px;
            border-radius: 5px;
            overflow-x: auto;
            font-size: 8pt;
        }
        
        .mermaid-diagram {
            margin: 20px 0;
            text-align: center;
            page-break-inside: avoid;
        }

        .mermaid-diagram-fallback {
            margin: 20px 0;
            padding: 10px;
            border: 1px solid #ddd;
            background: #f9f9f9;
            page-break-inside: avoid;
        }

        .mermaid-diagram-fallback pre {
            white-space: pre-wrap;
            word-break: break-word;
            font-size: 9pt;
        }
        
        .mermaid-diagram img {
            max-width: 100%;
            height: auto;
            display: block;
            margin: 0 auto;
        }
        
        tr {
            page-break-inside: avoid;
        }
        """

    # Step 5: Generate PDF
    html_full = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>{css}</style>
</head>
<body>
{html_body}
</body>
</html>"""

    HTML(string=html_full).write_pdf(output_pdf)
    return output_pdf

#====================for testcase generation==================================================

def convert_markdown_to_pdf_tc(md_path, pdf_path):
    """
    Convert Markdown to PDF using pdfkit (for Test Case workflow only).
    Alternative to WeasyPrint to avoid system dependency issues.
    """
    try:
        import markdown2
        import pdfkit
    except ImportError:
        print("⚠️  Missing dependencies for PDF generation.")
        print("   Install with: pip install markdown2 pdfkit")
        print("   Also install wkhtmltopdf:")
        print("   - Ubuntu/Debian: sudo apt-get install wkhtmltopdf")
        print("   - macOS: brew install wkhtmltopdf")
        print("   - Windows: Download from https://wkhtmltopdf.org/downloads.html")
        raise
    
    # Read markdown
    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Convert to HTML
    html = markdown2.markdown(md_content, extras=['tables', 'fenced-code-blocks'])
    
    # Add basic styling
    html_with_style = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            body {{ 
                font-family: Arial, sans-serif; 
                margin: 40px; 
                line-height: 1.6;
            }}
            h1 {{ 
                color: #333; 
                border-bottom: 2px solid #0066cc;
                padding-bottom: 10px;
            }}
            h2 {{ 
                color: #0066cc; 
                margin-top: 30px;
            }}
            h3 {{ 
                color: #666; 
            }}
            table {{ 
                border-collapse: collapse; 
                width: 100%; 
                margin: 20px 0;
            }}
            th, td {{ 
                border: 1px solid #ddd; 
                padding: 12px; 
                text-align: left; 
            }}
            th {{ 
                background-color: #0066cc; 
                color: white;
                font-weight: bold;
            }}
            tr:nth-child(even) {{
                background-color: #f9f9f9;
            }}
            code {{ 
                background-color: #f4f4f4; 
                padding: 2px 6px; 
                border-radius: 3px;
                font-family: 'Courier New', monospace;
            }}
            pre {{ 
                background-color: #f4f4f4; 
                padding: 15px; 
                border-radius: 5px;
                border-left: 4px solid #0066cc;
                overflow-x: auto;
            }}
            blockquote {{
                border-left: 4px solid #ddd;
                padding-left: 15px;
                color: #666;
                margin: 20px 0;
            }}
        </style>
    </head>
    <body>
        {html}
    </body>
    </html>
    """
    
    # Convert to PDF
    try:
        pdfkit.from_string(html_with_style, pdf_path)
    except Exception as e:
        print(f"❌ PDF conversion failed: {e}")
        print("   Make sure wkhtmltopdf is installed on your system")
        raise


def process_markdown(md_file: str):
    with open(md_file, "r", encoding="utf-8") as f:
        raw = f.read()

    cleaned = clean_markdown(raw)

    with open(md_file, "w", encoding="utf-8") as f:
        f.write(cleaned)

    output_pdf = md_file.replace(".txt", ".pdf")
    convert_markdown_to_pdf(md_file, output_pdf)
    print(f"PDF saved at {output_pdf}")


def _candidate_model_names_for_lookup(model_name: str):
    """Return possible model-name variants for SDK/AI Core lookup."""
    name = str(model_name or "").strip()
    if not name:
        return []

    candidates = [name]

    # If short alias is used, add anthropic prefix variant.
    if name.startswith("claude-"):
        candidates.append(f"anthropic--{name}")

    def add_if_new(value: str):
        if value and value not in candidates:
            candidates.append(value)

    # Convert old style: anthropic--claude-opus-4-5 -> anthropic--claude-4.5-opus
    for raw in list(candidates):
        m = re.match(r"^(anthropic--)?claude-(opus|sonnet|haiku)-(\d)-(\d)$", raw)
        if m:
            prefix = m.group(1) or ""
            tier = m.group(2)
            major = m.group(3)
            minor = m.group(4)
            add_if_new(f"{prefix}claude-{major}.{minor}-{tier}")

    # Convert new style: anthropic--claude-4.5-opus -> anthropic--claude-opus-4-5
    for raw in list(candidates):
        m = re.match(r"^(anthropic--)?claude-(\d)\.(\d)-(opus|sonnet|haiku)$", raw)
        if m:
            prefix = m.group(1) or ""
            major = m.group(2)
            minor = m.group(3)
            tier = m.group(4)
            add_if_new(f"{prefix}claude-{tier}-{major}-{minor}")

    # Include short alias if prefixed name was provided.
    if name.startswith("anthropic--claude-"):
        add_if_new(name.replace("anthropic--", "", 1))

    return candidates


def _dep_value(dep, key):
    if isinstance(dep, dict):
        return dep.get(key)
    return getattr(dep, key, None)


def _resolve_deployment_for_model(proxy_client, model_name: str):
    """Find deployment by deployment_id or model_name variants."""

    deployments = []
    try:
        deployments = list(proxy_client.get_deployments())
    except Exception:
        deployments = list(getattr(proxy_client, "deployments", []) or [])

    if not deployments:
        return None

    lookup = str(model_name or "").strip().lower()
    if lookup:
        for dep in deployments:
            dep_id = (
                _dep_value(dep, "deployment_id") or _dep_value(dep, "id") or ""
            ).lower()
            if dep_id and dep_id == lookup:
                return dep

    for candidate in _candidate_model_names_for_lookup(model_name):
        c = candidate.lower()
        for dep in deployments:
            dep_model_name = (_dep_value(dep, "model_name") or "").lower()
            if dep_model_name and dep_model_name == c:
                return dep

    return None


class _AICoreUsageShim:
    def __init__(self, prompt_tokens=0, completion_tokens=0, total_tokens=0):
        self.prompt_tokens = int(prompt_tokens or 0)
        self.completion_tokens = int(completion_tokens or 0)
        self.total_tokens = int(
            total_tokens or (self.prompt_tokens + self.completion_tokens)
        )


class _AICoreResponseShim:
    def __init__(self, content: str, usage_metadata: dict, usage: object):
        self.content = content
        self.usage_metadata = usage_metadata
        self.usage = usage


class _AICoreDeploymentLLMWrapper:
    """Fallback wrapper for deployments unsupported by SDK model registry."""

    def __init__(
        self, deployment_id: str, model_name: str, max_tokens: int, temperature: float
    ):
        self.deployment_id = deployment_id
        self.model_name = model_name
        self.max_tokens = int(max_tokens or 4096)
        self.temperature = temperature

    def _normalize_messages(self, messages):
        if isinstance(messages, str):
            return [{"role": "user", "content": messages}]

        normalized = []
        for item in messages or []:
            if isinstance(item, dict):
                role = str(item.get("role") or "user").lower()
                content = item.get("content", "")
            else:
                role = str(getattr(item, "type", "user")).lower()
                content = getattr(item, "content", "")

            if role in ("human", "user"):
                role = "user"
            elif role in ("ai", "assistant"):
                role = "assistant"
            elif role != "system":
                role = "user"

            if isinstance(content, list):
                parts = []
                for part in content:
                    if isinstance(part, dict):
                        parts.append(str(part.get("text", part)))
                    else:
                        parts.append(str(part))
                content = "\n".join(parts)

            normalized.append({"role": role, "content": str(content or "")})

        if not normalized:
            normalized = [{"role": "user", "content": ""}]

        return normalized

    def _invoke_openai_chat(self, normalized_messages):
        from gen_ai_hub.proxy.native.openai import chat as native_chat

        attempts = []

        kwargs = {
            "deployment_id": self.deployment_id,
            "messages": normalized_messages,
            "max_completion_tokens": self.max_tokens,
        }
        if self.temperature is not None:
            kwargs["temperature"] = self.temperature
        attempts.append(kwargs)

        kwargs_no_temp = dict(kwargs)
        kwargs_no_temp.pop("temperature", None)
        attempts.append(kwargs_no_temp)

        kwargs_max_tokens = dict(kwargs)
        kwargs_max_tokens.pop("max_completion_tokens", None)
        kwargs_max_tokens["max_tokens"] = self.max_tokens
        attempts.append(kwargs_max_tokens)

        kwargs_max_tokens_no_temp = dict(kwargs_max_tokens)
        kwargs_max_tokens_no_temp.pop("temperature", None)
        attempts.append(kwargs_max_tokens_no_temp)

        last_error = None
        response = None
        for attempt in attempts:
            try:
                response = native_chat.completions.create(**attempt)
                break
            except Exception as e:
                last_error = e

        if response is None:
            raise last_error

        return response

    def _invoke_bedrock_converse(self, normalized_messages):
        from botocore.config import Config
        from gen_ai_hub.proxy.native.amazon.clients import Session

        bedrock_messages = []
        system_entries = []

        for msg in normalized_messages:
            role = msg.get("role", "user")
            content_text = str(msg.get("content", ""))

            if role == "system":
                if content_text.strip():
                    system_entries.append({"text": content_text})
                continue

            if role not in ("user", "assistant"):
                role = "user"

            bedrock_messages.append(
                {
                    "role": role,
                    "content": [{"text": content_text}],
                }
            )

        if not bedrock_messages:
            bedrock_messages = [{"role": "user", "content": [{"text": ""}]}]

        converse_kwargs = {
            "messages": bedrock_messages,
            "inferenceConfig": {
                "maxTokens": self.max_tokens,
            },
        }
        if self.temperature is not None:
            converse_kwargs["inferenceConfig"]["temperature"] = self.temperature
        if system_entries:
            converse_kwargs["system"] = system_entries

        read_timeout = int(os.getenv("AICORE_BEDROCK_READ_TIMEOUT", "900"))
        connect_timeout = int(os.getenv("AICORE_BEDROCK_CONNECT_TIMEOUT", "30"))
        max_attempts = int(os.getenv("AICORE_BEDROCK_MAX_ATTEMPTS", "3"))

        bedrock_config = Config(
            read_timeout=read_timeout,
            connect_timeout=connect_timeout,
            retries={"max_attempts": max_attempts, "mode": "standard"},
        )

        client = Session().client(
            deployment_id=self.deployment_id,
            config=bedrock_config,
        )
        return client.converse(**converse_kwargs)

    @staticmethod
    def _extract_content_from_openai_response(response):
        content = ""
        try:
            if response.choices:
                content = response.choices[0].message.content or ""
        except Exception:
            content = ""

        if isinstance(content, list):
            content_parts = []
            for part in content:
                if isinstance(part, dict):
                    content_parts.append(str(part.get("text", part)))
                else:
                    content_parts.append(str(part))
            content = "\n".join(content_parts)

        return str(content)

    @staticmethod
    def _extract_usage_from_openai_response(response):
        usage = getattr(response, "usage", None)
        prompt_tokens = getattr(usage, "prompt_tokens", None)
        if prompt_tokens is None:
            prompt_tokens = getattr(usage, "input_tokens", 0)

        completion_tokens = getattr(usage, "completion_tokens", None)
        if completion_tokens is None:
            completion_tokens = getattr(usage, "output_tokens", 0)

        total_tokens = getattr(usage, "total_tokens", None)
        if total_tokens is None:
            total_tokens = (prompt_tokens or 0) + (completion_tokens or 0)

        return {
            "input_tokens": int(prompt_tokens or 0),
            "output_tokens": int(completion_tokens or 0),
            "total_tokens": int(total_tokens or 0),
        }

    @staticmethod
    def _extract_content_from_bedrock_response(response):
        output = response.get("output", {}) if isinstance(response, dict) else {}
        message = output.get("message", {}) if isinstance(output, dict) else {}
        content_blocks = message.get("content", []) if isinstance(message, dict) else []

        content_parts = []
        for block in content_blocks:
            if isinstance(block, dict):
                if "text" in block:
                    content_parts.append(str(block.get("text", "")))
                else:
                    content_parts.append(str(block))
            else:
                content_parts.append(str(block))

        return "\n".join(content_parts).strip()

    @staticmethod
    def _extract_usage_from_bedrock_response(response):
        usage = response.get("usage", {}) if isinstance(response, dict) else {}
        input_tokens = int(usage.get("inputTokens", 0) or 0)
        output_tokens = int(usage.get("outputTokens", 0) or 0)
        total_tokens = int(usage.get("totalTokens", input_tokens + output_tokens) or 0)

        return {
            "input_tokens": input_tokens,
            "output_tokens": output_tokens,
            "total_tokens": total_tokens,
        }

    def invoke(self, messages):
        normalized_messages = self._normalize_messages(messages)

        model_lower = (self.model_name or "").lower()
        prefer_bedrock = model_lower.startswith(
            "anthropic--"
        ) or model_lower.startswith("amazon--")

        response_kind = None
        response_obj = None
        openai_error = None

        if not prefer_bedrock:
            try:
                response_obj = self._invoke_openai_chat(normalized_messages)
                response_kind = "openai"
            except Exception as e:
                openai_error = e

        if response_obj is None:
            try:
                response_obj = self._invoke_bedrock_converse(normalized_messages)
                response_kind = "bedrock"
            except Exception:
                if openai_error is not None:
                    raise openai_error
                raise

        if response_kind == "openai":
            content = self._extract_content_from_openai_response(response_obj)
            usage_metadata = self._extract_usage_from_openai_response(response_obj)
        else:
            content = self._extract_content_from_bedrock_response(response_obj)
            usage_metadata = self._extract_usage_from_bedrock_response(response_obj)

        usage_obj = _AICoreUsageShim(
            prompt_tokens=usage_metadata["input_tokens"],
            completion_tokens=usage_metadata["output_tokens"],
            total_tokens=usage_metadata["total_tokens"],
        )

        return _AICoreResponseShim(
            content=str(content), usage_metadata=usage_metadata, usage=usage_obj
        )


# -------------------------------------------------------------------
# LLM INITIALIZATION — UNIFIED SAP AI CORE + DEBUG MODE
# -------------------------------------------------------------------


def initialize_llm_model(model_name="gpt-5-nano", max_tokens=None, temperature=None):
    """
    Initialize and return an LLM model using SAP AI Core GenAI Hub.

    Configuration is loaded from model_config.json. If max_tokens or temperature
    are not provided, they will be loaded from the configuration file.

    Args:
        model_name: Model identifier (must exist in model_config.json)
        max_tokens: Override max tokens (optional, uses config default if None)
        temperature: Override temperature (optional, uses config default if None)

    Modes:
    - If DEBUG=true in .env → Use Gemini API (for testing without SAP AI Core)
    - Otherwise → Use SAP AI Core (production)

    Supported AI Core models (see model_config.json for full list):
    - gpt-5-nano: OpenAI GPT-5 Nano model
    - gemini-2.5-pro: Google Gemini 2.5 Pro model
    - claude-3-5-sonnet: Anthropic Claude 3.5 Sonnet
    - And more...

    Both models are deployed and accessible via the gen_ai_hub.proxy.langchain.init_llm function.
    """

    # Load model configuration from model_config.json
    import json

    config_path = os.path.join(os.path.dirname(__file__), "model_config.json")

    try:
        with open(config_path, "r") as f:
            model_configs = json.load(f)
    except FileNotFoundError:
        print(f"⚠️ WARNING: model_config.json not found at {config_path}")
        print("⚠️ Using default parameters")
        model_configs = {}
    except json.JSONDecodeError as e:
        print(f"⚠️ WARNING: Invalid JSON in model_config.json: {e}")
        print("⚠️ Using default parameters")
        model_configs = {}

    # Get model configuration (with fallback defaults)
    if model_name in model_configs:
        config = model_configs[model_name]
        # Use config defaults if parameters not provided
        max_tokens = (
            max_tokens if max_tokens is not None else config.get("max_tokens", 4096)
        )
        temperature = (
            temperature if temperature is not None else config.get("temperature", 0.7)
        )

        print(f"🤖 Initializing model: {model_name}")
        print(f"   Max tokens: {max_tokens}")
        print(f"   Temperature: {temperature}")
        print(f"   Description: {config.get('description', 'N/A')}")
    else:
        print(f"⚠️ WARNING: Model '{model_name}' not found in model_config.json")
        print("⚠️ Using default parameters")
        max_tokens = max_tokens if max_tokens is not None else 4096
        temperature = temperature if temperature is not None else 0.7
        print(f"🤖 Initializing model: {model_name}")
        print(f"   Max tokens: {max_tokens}")
        print(f"   Temperature: {temperature}")

    # Check if DEBUG mode is enabled
    debug_mode = os.getenv("DEBUG", "false").lower() == "true"

    if debug_mode:
        # DEBUG MODE: Use Gemini API directly (no SAP AI Core)
        gemini_api_key = os.getenv("GEMINI_API_KEY")

        if not gemini_api_key:
            print("⚠️ WARNING: DEBUG=true but GEMINI_API_KEY not found in .env")
            print("⚠️ Falling back to SAP AI Core")
        else:
            print("🔧 DEBUG MODE: Using Gemini API directly")
            return _initialize_gemini_llm(
                gemini_api_key, model_name, max_tokens, temperature
            )

    # PRODUCTION MODE: Use SAP AI Core with init_llm
    # This approach works for all models deployed in SAP AI Core
    print(f"🚀 Connecting to SAP AI Core...")

    try:
        from gen_ai_hub.proxy.langchain.init_models import init_llm
        from gen_ai_hub.proxy.core.proxy_clients import get_proxy_client

        # Suppress deprecation warnings from VertexAI SDK
        import warnings

        warnings.filterwarnings("ignore", category=UserWarning, module="vertexai")
        warnings.filterwarnings("ignore", message=".*genaihub_client.*")

        # Get proxy client for SAP AI Core
        proxy_client = get_proxy_client("gen-ai-hub")

        # Try the requested model and known naming variants first.
        # This covers alias differences such as:
        # - anthropic--claude-opus-4-5
        # - anthropic--claude-4.5-opus
        last_registry_error = None
        for candidate_name in _candidate_model_names_for_lookup(model_name):
            try:
                llm = init_llm(
                    model_name=candidate_name,
                    proxy_client=proxy_client,
                    max_tokens=max_tokens,
                    temperature=temperature,
                )
                print(f"✅ Model {candidate_name} initialized successfully!")
                return llm
            except KeyError as e:
                last_registry_error = e

        # Fallback: If model is deployed in AI Core but unavailable in SDK model registry,
        # invoke it via deployment_id using the native OpenAI-compatible proxy client.
        deployment = _resolve_deployment_for_model(proxy_client, model_name)
        if deployment:
            deployment_id = _dep_value(deployment, "deployment_id") or _dep_value(
                deployment, "id"
            )
            deployment_model_name = _dep_value(deployment, "model_name") or model_name
            if deployment_id:
                print(
                    "⚠️ Model is not available in SDK registry; "
                    f"falling back to deployment_id={deployment_id} "
                    f"(model_name={deployment_model_name})"
                )
                return _AICoreDeploymentLLMWrapper(
                    deployment_id=deployment_id,
                    model_name=deployment_model_name,
                    max_tokens=max_tokens,
                    temperature=temperature,
                )

        if last_registry_error is not None:
            print(f"❌ Error: {last_registry_error}")
            print(f"⚠️ Model '{model_name}' not available in SAP AI Core")
            print("⚠️ No matching deployment found for deployment-id fallback")
            print(
                "⚠️ Make sure the model/deployment is running and "
                "sap-ai-sdk-gen is installed"
            )
            raise last_registry_error

        raise RuntimeError(
            f"Unable to initialize model '{model_name}' via SDK or deployment fallback"
        )

    except Exception as e:
        print(f"❌ Error initializing model: {e}")
        import traceback

        traceback.print_exc()
        raise


# def _initialize_gemini_llm(api_key, model_name="gemini-2.5-pro", max_tokens=4096, temperature=0.7):
#     """
#     Initialize Gemini API for DEBUG mode using native google-generativeai library.
#     This bypasses SAP AI Core and connects directly to Google's Gemini API.

#     Note: max_tokens parameter is ignored for direct Gemini API calls.
#     """
#     try:
#         import google.generativeai as genai

#         # Configure Gemini with API key
#         genai.configure(api_key=api_key)

#         # Wrap to match LangChain interface
#         class GeminiLLMWrapper:
#             def __init__(self, model_name, temperature):
#                 self.model = genai.GenerativeModel(model_name)
#                 self.temperature = temperature

#             def invoke(self, messages):
#                 # Convert LangChain messages to Gemini format
#                 formatted_messages = []
#                 system_instruction = None

#                 for m in messages:
#                     if m.type == "system":
#                         # Gemini handles system prompts separately
#                         system_instruction = m.content
#                     elif m.type == "human" or m.type == "user":
#                         formatted_messages.append({
#                             "role": "user",
#                             "parts": [m.content]
#                         })
#                     elif m.type == "ai" or m.type == "assistant":
#                         formatted_messages.append({
#                             "role": "model",
#                             "parts": [m.content]
#                         })

#                 # Combine system instruction with first user message if present
#                 if system_instruction and formatted_messages:
#                     if formatted_messages[0]["role"] == "user":
#                         formatted_messages[0]["parts"][0] = (
#                             f"{system_instruction}\n\n{formatted_messages[0]['parts'][0]}"
#                         )

#                 # Generate response
#                 generation_config = genai.types.GenerationConfig(
#                     temperature=self.temperature,
#                 )

#                 # For single-turn conversation
#                 if len(formatted_messages) == 1:
#                     response = self.model.generate_content(
#                         formatted_messages[0]["parts"][0],
#                         generation_config=generation_config
#                     )
#                 else:
#                     # For multi-turn conversation
#                     chat = self.model.start_chat(history=formatted_messages[:-1])
#                     response = chat.send_message(
#                         formatted_messages[-1]["parts"][0],
#                         generation_config=generation_config
#                     )

#                 # Create usage mock (Gemini doesn't provide OpenAI-style usage)
#                 class UsageMock:
#                     def __init__(self, prompt_tokens=0, completion_tokens=0):
#                         self.prompt_tokens = prompt_tokens
#                         self.completion_tokens = completion_tokens
#                         self.total_tokens = prompt_tokens + completion_tokens

#                 # Estimate token usage (rough approximation)
#                 prompt_text = "".join([msg["parts"][0] for msg in formatted_messages])
#                 estimated_input = len(prompt_text) // 4  # Rough estimate: 4 chars per token
#                 estimated_output = len(response.text) // 4

#                 usage = UsageMock(
#                     prompt_tokens=estimated_input,
#                     completion_tokens=estimated_output
#                 )

#                 # Create response wrapper to match LangChain interface
#                 class ResponseWrapper:
#                     def __init__(self, content, usage):
#                         self.content = content
#                         self.usage = usage

#                 return ResponseWrapper(response.text, usage)

#         return GeminiLLMWrapper(model_name, temperature)

#     except ImportError:
#         print("⚠️ ERROR: google-generativeai not installed")
#         print("⚠️ Install with: pip install google-generativeai")
#         raise


# -------------------------------------------------------------------
# OPTIONAL TOOLS
# -------------------------------------------------------------------


def list_available_deployments():
    try:
        from gen_ai_hub.proxy.core.proxy_clients import get_proxy_client

        client = get_proxy_client("gen-ai-hub")
        deployments = client.get_deployments()

        print("\n📋 Available Deployments")
        for d in deployments:
            print(f"- {d.model_name} → {d.deployment_id}")

    except Exception as e:
        print("Error listing deployments:", e)


def query_llm_model(llm, message):
    print(f"→ Sending: {message}")
    resp = llm.invoke(message)
    print("← Response:", resp.content)
    return resp.content

# =====================================================================
# DOCUMENT READING UTILITY
# =====================================================================

def read_document_content(file_path):
    """
    Read content from various document formats.
    Supports: PDF, TXT, MD, DOCX
    
    Args:
        file_path: Path to document file
        
    Returns:
        str: Extracted text content
    """
    import os
    from pathlib import Path
    
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Document not found: {file_path}")
    
    file_ext = Path(file_path).suffix.lower()
    
    print(f"📄 Reading document: {file_path} (type: {file_ext})")
    
    # Text files (TXT, MD)
    if file_ext in ['.txt', '.md']:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            print(f"✅ Read {len(content)} characters from text file")
            return content
        except Exception as e:
            print(f"⚠️  Error reading text file, trying with different encoding...")
            with open(file_path, 'r', encoding='latin-1') as f:
                content = f.read()
            return content
    
    # PDF files
    elif file_ext == '.pdf':
        try:
            import PyPDF2
            content = []
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                num_pages = len(pdf_reader.pages)
                print(f"📖 PDF has {num_pages} pages")
                
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text:
                        content.append(text)
            
            full_content = '\n\n'.join(content)
            print(f"✅ Extracted {len(full_content)} characters from PDF")
            return full_content
            
        except ImportError:
            raise ImportError("PyPDF2 is required to read PDF files. Install it with: pip install PyPDF2")
        except Exception as e:
            print(f"❌ Error reading PDF: {e}")
            raise
    
    # DOCX files
    elif file_ext == '.docx':
        try:
            import docx
            doc = docx.Document(file_path)
            content = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            content.append(cell.text)
            
            full_content = '\n\n'.join(content)
            print(f"✅ Extracted {len(full_content)} characters from DOCX")
            return full_content
            
        except ImportError:
            raise ImportError("python-docx is required to read DOCX files. Install it with: pip install python-docx")
        except Exception as e:
            print(f"❌ Error reading DOCX: {e}")
            raise
    
    else:
        raise ValueError(f"Unsupported file format: {file_ext}. Supported formats: .txt, .md, .pdf, .docx")